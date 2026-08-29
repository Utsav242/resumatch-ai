import io
from typing import Any, Dict, Tuple

import docx
from fastapi import HTTPException, status

from app.core.logging import logger
from app.services.parsing.text_normalizer import normalize_text


def parse_docx(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """Parse a DOCX document using python-docx and extract text.

    Returns:
        A tuple of (normalized_text, metadata_dict)
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error("Failed to open DOCX document", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to parse DOCX document. The file might be corrupted or invalid.",
        ) from e

    text_parts = []
    metadata = {"page_count": None}  # DOCX doesn't easily store pages without rendering

    try:
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)

        # Simple extraction from tables as well to capture work history/details in grid layouts
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
    except Exception as e:
        logger.error("Error occurred while extracting text from DOCX", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Failed to extract text from the DOCX file.",
        ) from e

    raw_extracted_text = "\n".join(text_parts)
    normalized_text = normalize_text(raw_extracted_text)

    if not normalized_text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="The uploaded DOCX file does not contain any readable text.",
        )

    return normalized_text, metadata
